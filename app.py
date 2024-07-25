import json
import os
import sys
import boto3
import streamlit as st
import docx
import uuid
from PyPDF2 import PdfReader  # Import PyPDF2 for PDF parsing

# We will be using Titan Embeddings Model to generate Embedding
from langchain_community.embeddings import BedrockEmbeddings
from langchain_aws import BedrockLLM  # Use the updated import

# Data Ingestion
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFDirectoryLoader

# Vector Embedding and Vector Store
import faiss
from langchain_community.vectorstores import FAISS

# Set the region
region_name = "us-east-1"

# Bedrock Clients
bedrock = boto3.client(service_name="bedrock-runtime", region_name=region_name)

# Initialize AWS clients for transcription and connect
transcribe = boto3.client('transcribe', region_name=region_name)
connect = boto3.client('connect', region_name=region_name)

# Data ingestion
def parse_resume(file):
    if file.name.endswith('.pdf'):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError('Unsupported file format')

    splitter = RecursiveCharacterTextSplitter()
    chunks = splitter.split_text(text)
    return chunks

def embed_documents(docs):
    bedrock_embeddings = BedrockEmbeddings(model_id="amazon.titan-embed-text-v1", client=bedrock)
    return [bedrock_embeddings.embed_query(doc) for doc in docs]

def get_vector_store(embeddings):
    index = faiss.IndexFlatL2(len(embeddings[0]))
    index.add(np.array(embeddings).astype('float32'))
    return index

def get_llama3_llm():
    # Create the Llama 3 Model
    return BedrockLLM(model_id="meta.llama3-70b-instruct-v1:0", client=bedrock, model_kwargs={'max_gen_len':128})

def truncate_text(text, max_tokens=500):
    """Truncate text to ensure it doesn't exceed the token limit."""
    words = text.split()
    if len(words) > max_tokens:
        return ' '.join(words[:max_tokens]) + '...'
    return text

def clean_standard_questions(questions):
    cleaned_questions = []
    for question in questions:
        cleaned_question = ' '.join(question.split())  # Remove extra spaces and newlines
        cleaned_questions.append(cleaned_question)
    return cleaned_questions

# Prompt template to generate new questions
def create_prompt(context, standard_questions, transcript):
    truncated_context = truncate_text(context)
    cleaned_questions = clean_standard_questions(standard_questions)
    truncated_standard_questions = truncate_text(' '.join(cleaned_questions), max_tokens=200)
    truncated_transcript = truncate_text(transcript, max_tokens=200)
    
    return f"""
    Based on the following resume content, standard interview questions, and the interviewee's recent response, generate a concise follow-up interview question (1 or 2 sentences) tailored to the interviewee.

    Resume Content:
    {truncated_context}

    Standard Interview Questions:
    {truncated_standard_questions}

    Interviewee's Response:
    {truncated_transcript}

    New Follow-Up Question:
    """

def generate_question(llm, context, standard_questions, transcript):
    prompt = create_prompt(context, standard_questions, transcript)
    try:
        response = llm.invoke(prompt)
        return response.strip()  # Assuming the response is plain text
    except Exception as e:
        st.error(f"Error generating question: {e}")
        return "An error occurred while generating the question."

def start_transcription_streaming(resume_text, standard_questions):
    import asyncio
    from websockets import connect

    async def transcribe_streaming():
        async with connect(
            f"wss://transcribestreaming.{region_name}.amazonaws.com:8443/stream-transcription-websocket?media-encoding=pcm&sample-rate=16000"
        ) as ws:
            async def send_audio():
                with open("audio.wav", "rb") as f:
                    while chunk := f.read(4096):
                        await ws.send(chunk)
                await ws.send(json.dumps({'event': 'AudioEvent', 'AudioData': None}))  # Send the end signal

            async def receive_transcript():
                async for message in ws:
                    event = json.loads(message)
                    if 'Transcript' in event:
                        transcript = event['Transcript']['Results'][0]['Alternatives'][0]['Transcript']
                        llm = get_llama3_llm()
                        new_question = generate_question(llm, resume_text, standard_questions, transcript)
                        st.write("Generated Follow-Up Question: ", new_question)

            await asyncio.gather(send_audio(), receive_transcript())

    asyncio.run(transcribe_streaming())

def parse_prompts_from_file(file):
    if file.name.endswith('.pdf'):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        prompts = text.strip().split('\n')
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        prompts = [para.text for para in doc.paragraphs]
    else:
        raise ValueError('Unsupported file format')
    return prompts

def start_session(phone_number):
    instance_id = os.environ.get('CONNECT_INSTANCE_ID', 'your_instance_id')
    contact_flow_id = os.environ.get('CONNECT_CONTACT_FLOW_ID', 'your_contact_flow_id')
    response = connect.start_outbound_voice_contact(
        DestinationPhoneNumber=phone_number,
        ContactFlowId=contact_flow_id,
        InstanceId=instance_id,
        SourcePhoneNumber=os.environ.get('SOURCE_PHONE_NUMBER', 'your_source_phone_number')
    )
    return response

# Initialize session state for standard questions
if "standard_questions" not in st.session_state:
    st.session_state.standard_questions = []

# Streamlit UI
st.title("Interview Assistant App")

# Upload resume
st.header("Upload Resume")
uploaded_file = st.file_uploader("Choose a PDF or DOCX file", type=["pdf", "docx"])

resume_vector = None
resume_text = None

if uploaded_file is not None:
    resume_chunks = parse_resume(uploaded_file)
    resume_embeddings = embed_documents(resume_chunks)
    resume_vector = np.mean(resume_embeddings, axis=0)
    faiss_index = get_vector_store(resume_embeddings)
    resume_text = "\n".join(resume_chunks)
    st.success("Resume parsed and stored successfully.")

# Import standard prompts
st.header("Import Standard Prompts")
prompt_file = st.file_uploader("Choose a PDF or DOCX file with standard prompts", type=["pdf", "docx"])

if prompt_file is not None:
    prompts = parse_prompts_from_file(prompt_file)
    st.session_state.standard_questions = prompts
    st.success("Prompts imported successfully.")

# Generate interview questions
st.header("Generate Interview Questions")
if st.button("Generate Questions"):
    if uploaded_file and len(st.session_state.standard_questions) > 0:
        llm = get_llama3_llm()
        initial_question = generate_question(llm, resume_text, st.session_state.standard_questions, "")
        st.write("Initial Question: ", initial_question)
    else:
        st.warning("Please upload a resume and import standard interview questions first.")

# Real-time transcription
st.header("Real-time Transcription")
if st.button("Start Real-time Transcription"):
    if resume_vector is not None and resume_vector.any() and len(st.session_state.standard_questions) > 0:
        start_transcription_streaming(resume_text, st.session_state.standard_questions)
        st.success("Real-time transcription started.")
    else:
        st.warning("Please upload a resume and import standard interview questions first.")

# Start communication session
st.header("Start Communication Session")
phone_number = st.text_input("Phone Number")
if st.button("Start Session"):
    start_session(phone_number)
    st.success("Session started.")

if __name__ == '__main__':
    if len(sys.argv) > 1 and sys.argv[1] == 'run':
        st.run()
    else:
        st.warning("Run the script with 'streamlit run app.py' to start the Streamlit server.")
